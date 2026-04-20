"""
Generate a DOCX list of all source documents used to build the IR corpus.

Source data:
- scripts/build_corpus.py (12 PDF-backed papers, file -> source mapping)
- ir_explorer/assets/default_corpus.json (15 unique sources, 80 docs)
- Initial corpus commit ea78f1f (3 seed sources without PDFs)
"""

import json
import os
from collections import defaultdict

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


PROJECT_ROOT = os.path.join(os.path.dirname(__file__), "..")
CORPUS_PATH = os.path.join(PROJECT_ROOT, "ir_explorer", "assets", "default_corpus.json")
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "Corpus_Sources.docx")


# Manually curated bibliographic records. Titles, venues, years were
# cross-checked against the PDF first-page content and public records.
SOURCES = [
    # ---- Artificial Intelligence & Machine Learning ----
    {
        "theme": "Artificial Intelligence & Machine Learning",
        "citation": "Vaswani et al. (2017)",
        "title": "Attention Is All You Need",
        "authors": "Ashish Vaswani, Noam Shazeer, Niki Parmar, Jakob Uszkoreit, Llion Jones, Aidan N. Gomez, Lukasz Kaiser, Illia Polosukhin",
        "venue": "arXiv:1706.03762v7 (NeurIPS 2017)",
        "file": "1706.03762v7.pdf",
    },
    {
        "theme": "Artificial Intelligence & Machine Learning",
        "citation": "Devlin et al. (2019)",
        "title": "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding",
        "authors": "Jacob Devlin, Ming-Wei Chang, Kenton Lee, Kristina Toutanova",
        "venue": "arXiv:1810.04805v2 (NAACL 2019)",
        "file": "1810.04805v2.pdf",
    },
    {
        "theme": "Artificial Intelligence & Machine Learning",
        "citation": "LeCun, Bengio & Hinton (2015)",
        "title": "Deep Learning",
        "authors": "Yann LeCun, Yoshua Bengio, Geoffrey Hinton",
        "venue": "Nature, 521(7553), 436-444",
        "file": "Deep Learning Review (Nature) - Yann LeCun, Yoshua Bengio, Geoffrey Hinton.pdf",
    },
    {
        "theme": "Artificial Intelligence & Machine Learning",
        "citation": "Turing (1950)",
        "title": "Computing Machinery and Intelligence",
        "authors": "Alan M. Turing",
        "venue": "Mind, 59(236), 433-460",
        "file": "turing.pdf",
    },
    {
        "theme": "Artificial Intelligence & Machine Learning",
        "citation": "Hebb (1949)",
        "title": "The Organization of Behavior: A Neuropsychological Theory",
        "authors": "Donald O. Hebb",
        "venue": "Wiley, New York",
        "file": "Hebb_1949_The_Organization_of_Behavior.pdf",
    },
    # ---- Physics & Cosmology ----
    {
        "theme": "Physics & Cosmology",
        "citation": "Einstein (1905)",
        "title": "On the Electrodynamics of Moving Bodies (Zur Elektrodynamik bewegter Koerper)",
        "authors": "Albert Einstein",
        "venue": "Annalen der Physik, 17, 891-921",
        "file": "specrel.pdf",
    },
    {
        "theme": "Physics & Cosmology",
        "citation": "Hawking (1975)",
        "title": "Particle Creation by Black Holes",
        "authors": "Stephen W. Hawking",
        "venue": "Communications in Mathematical Physics, 43, 199-220",
        "file": "Download.pdf",
    },
    {
        "theme": "Physics & Cosmology",
        "citation": "Tegmark (2005)",
        "title": "Parallel Universes (The Multiverse Hierarchy)",
        "authors": "Max Tegmark",
        "venue": "In Science and Ultimate Reality (Cambridge Univ. Press); arXiv:astro-ph/0302131",
        "file": "(no PDF in repo - text seeded directly into default_corpus.json)",
    },
    {
        "theme": "Physics & Cosmology",
        "citation": "Guth (2007)",
        "title": "Eternal Inflation and Its Implications",
        "authors": "Alan H. Guth",
        "venue": "Journal of Physics A: Math. Theor., 40, 6811-6826; arXiv:hep-th/0702178",
        "file": "(no PDF in repo - text seeded directly into default_corpus.json)",
    },
    {
        "theme": "Physics & Cosmology",
        "citation": "Blackshaw & Franklin (2026)",
        "title": "Everettian Interpretations of Quantum Mechanics",
        "authors": "Nadia Blackshaw, Alexander Franklin",
        "venue": "Forthcoming in Comprehensive Philosophy of Science (Elsevier); PhilSci-Archive 28450",
        "file": "(no PDF in repo - text seeded directly into default_corpus.json)",
    },
    # ---- Information Theory ----
    {
        "theme": "Information Theory",
        "citation": "Shannon (1948)",
        "title": "A Mathematical Theory of Communication",
        "authors": "Claude E. Shannon",
        "venue": "Bell System Technical Journal, 27, 379-423 & 623-656",
        "file": "entropy.pdf",
    },
    # ---- Biology ----
    {
        "theme": "Biology",
        "citation": "Darwin (1859)",
        "title": "On the Origin of Species by Means of Natural Selection",
        "authors": "Charles Darwin",
        "venue": "John Murray, London (1st edition)",
        "file": "original.pdf",
    },
    {
        "theme": "Biology",
        "citation": "Watson & Crick (1953)",
        "title": "Molecular Structure of Nucleic Acids: A Structure for Deoxyribose Nucleic Acid",
        "authors": "James D. Watson, Francis H. C. Crick",
        "venue": "Nature, 171, 737-738",
        "file": "WatsonCrick1953.pdf",
    },
    # ---- Climate Science ----
    {
        "theme": "Climate Science",
        "citation": "Keeling (1960)",
        "title": "The Concentration and Isotopic Abundances of Carbon Dioxide in the Atmosphere",
        "authors": "Charles D. Keeling",
        "venue": "Tellus, 12(2), 200-203",
        "file": "6572a8eb567d8.pdf",
    },
    {
        "theme": "Climate Science",
        "citation": "IPCC AR6 (2021)",
        "title": "Climate Change 2021: The Physical Science Basis - Summary for Policymakers",
        "authors": "IPCC Working Group I (Masson-Delmotte et al., eds.)",
        "venue": "Sixth Assessment Report, Cambridge University Press",
        "file": "IPCC_AR6_WGI_SPM.pdf",
    },
]


def count_docs_per_source():
    with open(CORPUS_PATH, "r", encoding="utf-8") as f:
        corpus = json.load(f)
    counts = defaultdict(int)
    for d in corpus["documents"].values():
        counts[d.get("source", "?")] += 1
    return counts


def section_titles_per_source():
    with open(CORPUS_PATH, "r", encoding="utf-8") as f:
        corpus = json.load(f)
    by_src = defaultdict(list)
    for did, d in corpus["documents"].items():
        by_src[d.get("source", "?")].append((did, d.get("title", "")))
    return by_src


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def build_docx():
    counts = count_docs_per_source()
    sections = section_titles_per_source()

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("IR Explorer - Default Corpus Sources", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run(
        "This document lists every source paper used to build the default "
        "retrieval corpus for IR Explorer. The corpus contains "
    )
    total = sum(counts.values())
    intro.add_run(f"{total} section-level documents ").bold = True
    intro.add_run(f"drawn from {len(SOURCES)} distinct source works across five themes. ")
    intro.add_run(
        "Twelve works are extracted from PDF files in the repository root by "
    )
    r = intro.add_run("scripts/build_corpus.py")
    r.italic = True
    intro.add_run(
        "; three (Tegmark 2005, Guth 2007, Blackshaw & Franklin 2026) were "
        "seeded as hand-written summaries directly into "
    )
    r = intro.add_run("ir_explorer/assets/default_corpus.json")
    r.italic = True
    intro.add_run(".")

    # ---- Summary table ----
    add_heading(doc, "Summary Table", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "Citation", "Title", "Docs", "File in repo"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True

    for i, src in enumerate(SOURCES, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = src["citation"]
        row[2].text = src["title"]
        row[3].text = str(counts.get(src["citation"], 0))
        row[4].text = src["file"]

    # ---- Detailed entries grouped by theme ----
    add_heading(doc, "Detailed Entries by Theme", level=1)

    by_theme = defaultdict(list)
    for s in SOURCES:
        by_theme[s["theme"]].append(s)

    theme_order = [
        "Artificial Intelligence & Machine Learning",
        "Physics & Cosmology",
        "Information Theory",
        "Biology",
        "Climate Science",
    ]

    for theme in theme_order:
        add_heading(doc, theme, level=2)
        for src in by_theme[theme]:
            add_heading(doc, f"{src['citation']} - {src['title']}", level=3)

            p = doc.add_paragraph()
            p.add_run("Authors: ").bold = True
            p.add_run(src["authors"])

            p = doc.add_paragraph()
            p.add_run("Venue: ").bold = True
            p.add_run(src["venue"])

            p = doc.add_paragraph()
            p.add_run("File in repository: ").bold = True
            r = p.add_run(src["file"])
            r.italic = True

            p = doc.add_paragraph()
            p.add_run("Corpus sections: ").bold = True
            p.add_run(f"{counts.get(src['citation'], 0)} documents")

            # Show the extracted section titles/IDs
            secs = sections.get(src["citation"], [])
            if secs:
                p = doc.add_paragraph()
                p.add_run("Section IDs and titles in default_corpus.json:").bold = True
                for did, stitle in secs:
                    bullet = doc.add_paragraph(style="List Bullet")
                    bullet.add_run(f"{did}: ").bold = True
                    bullet.add_run(stitle if stitle else "(untitled)")

    # ---- Provenance note ----
    add_heading(doc, "Provenance & Verification Notes", level=1)
    notes = [
        "The 12 PDF-backed papers are defined in scripts/build_corpus.py (PAPERS list, lines 152-165). Each is processed by extract_sections() from ir_explorer/core/pdf_reader.py, cleaned of non-ASCII characters, filtered for non-content sections (references, bibliography, appendices), and capped at 5-6 section-level documents per paper.",
        "The 3 non-PDF sources (Tegmark 2005, Guth 2007, Blackshaw & Franklin 2026) were present in the initial corpus at commit ea78f1f ('IR Explorer') as hand-authored JSON entries and were preserved when scripts/build_corpus.py extended the corpus (commit a49d04e).",
        "PDF identities were verified by extracting first-page content: 'Download.pdf' begins with 'Particle Creation by Black Holes - S. W. Hawking ... Received April 12, 1975'; 'specrel.pdf' begins 'June 30, 1905 ... Maxwell's electrodynamics'; 'original.pdf' begins 'ORIGIN OF SPECIES'; 'turing.pdf' begins '1. The Imitation Game ... Can machines think?'; '6572a8eb567d8.pdf' begins 'La Jolla, California ... Manuscript received March 25, 1960' (Keeling); 'entropy.pdf' is 'Reprinted ... The Bell System Technical Journal, Vol. 27, pp. 379-423, 623-656, July, October, 1948' (Shannon).",
        "Blackshaw & Franklin (2026) was confirmed via web search against PhilSci-Archive entry 28450 ('Everettian Interpretations of Quantum Mechanics', dated 5 March 2026, forthcoming in Comprehensive Philosophy of Science, Elsevier).",
    ]
    for n in notes:
        p = doc.add_paragraph(n, style="List Bullet")

    doc.save(OUTPUT_PATH)
    print(f"Written: {OUTPUT_PATH}")
    print(f"Total sources: {len(SOURCES)} | Total docs: {total}")


if __name__ == "__main__":
    build_docx()
